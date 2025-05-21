from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Cover:
    def __init__(self, doc = Document()):
       self.doc = doc
    
    def run(self):
        self.header()
        self.title()
        self.footer()
    
    def header(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run()
        r.add_picture('assets/brasaooficialcolorido.png', width=Inches(1.25))
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(
            'MINISTÉRIO DA EDUCAÇÃO\n'
            'SECRETARIA DE EDUCAÇÃO PROFISSIONAL E TECNOLÓGICA\n'
            'INSTITUTO FEDERAL DE EDUCAÇÃO, CIÊNCIA E TECNOLOGIA DE PERNAMBUCO\n'
        )
        r = p.add_run('CAMPUS ')
        r.italic = True
        p.add_run('IGARASSU\nDIREÇÃO DE ENSINO')

    def title(self):
        for _ in range(10):
            self.doc.add_paragraph("")
            
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('PROJETO PEDAGÓGICO DO CURSO SUPERIOR DE TECNOLOGIA \n EM SISTEMAS PARA INTERNET')
        r.bold = True

    def footer(self):
        footer = self.doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Igarassu\n2025")

if __name__ == "__main__":
    c = Cover()
    c.run()
    c.doc.save('cover.docx')
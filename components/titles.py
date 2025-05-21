from docx import Document
from haggis.files.docx import list_number

class Titles:
    def __init__(self, doc = Document()):
        self.doc = doc

    def run(self):
        p1 = self.doc.add_paragraph('Text1', style='List Number')
        p2 = self.doc.add_paragraph('Text2', style='List Number')
        p3 = self.doc.add_paragraph('Text3', style='List Number')
        p4 = self.doc.add_paragraph('Text4', style='List Number')
        p5 = self.doc.add_paragraph('Text5', style='List Number')
        p6 = self.doc.add_paragraph('Text6', style='List Number')
        p7 = self.doc.add_paragraph('Text7', style='List Number')


        list_number(self.doc, p1, prev=None, level=0)
        list_number(self.doc, p2, prev=p1, level=1)
        list_number(self.doc, p3, prev=p2, level=1)
        list_number(self.doc, p4, prev=p3, level=0)
        list_number(self.doc, p5, prev=p4, level=1)
        list_number(self.doc, p6, prev=p5, level=2)
        list_number(self.doc, p7, prev=p6, level=2)
    
    

if __name__ == "__main__":
    Titles().run().save('titles.docx')